import os
from sqlalchemy import *
from docx import Document
from sqlalchemy.orm import sessionmaker
from db_connection.get_set.getter import Getter
from db_connection.get_set.setter import Setter


class DBClass:
    def __init__(self):
        self.engine = create_engine("postgresql+psycopg2://postgres:Chronos386@localhost/face_ident_db")
        self.Session = sessionmaker(bind=self.engine)
        self.session = self.Session()
        self.getter = Getter(self.session)
        self.setter = Setter(self.session)

    # Функция для создания документа Word с таблицами из базы данных
    def create_word_document(self):
        doc = Document()
        tables = [
            ("Metadata", self.getter.getAllMetadata()),
            ("Person", self.getter.getAllPersons()),
            ("Subject", self.getter.getAllSubjects()),
            ("Student's group", self.getter.getAllStudentsGroup()),
            ("Student", self.getter.getAllStudents()),
            ("Teacher", self.getter.getAllTeachers()),
            ("Class", self.getter.getAllClasses()),
            ("Class-teacher", self.getter.getAllClassTeacher()),
            ("Class-group", self.getter.getAllClassGroup()),
            ("Subject-teacher", self.getter.getAllSubjectTeacher()),
            ("Attendance", self.getter.getAllAttendance()),
        ]

        for table_name, data in tables:
            doc.add_heading(table_name, level=1)
            if len(data) != 0:
                my_table = doc.add_table(rows=1, cols=len(data[0].__table__.columns), style='Table Grid')

                for i, my_column in enumerate(data[0].__table__.columns):
                    my_table.cell(0, i).text = my_column.key

                for row_data in data:
                    row = my_table.add_row().cells
                    for i, my_column in enumerate(row_data.__table__.columns):
                        cell_value = str(getattr(row_data, my_column.key))
                        row[i].text = cell_value

                for _ in range(2):
                    doc.add_paragraph("")
            else:
                doc.add_paragraph("Таблица пуста")

        desktop_path = os.path.expanduser("~/Desktop")
        doc_path = os.path.join(desktop_path, "tables.docx")
        doc.save(doc_path)
